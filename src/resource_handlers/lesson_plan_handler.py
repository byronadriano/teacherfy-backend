# src/resource_handlers/lesson_plan_handler.py - CLEANED VERSION
import os
import logging
import docx
from typing import Dict, Any, List
from .base_handler import BaseResourceHandler

logger = logging.getLogger(__name__)

class LessonPlanHandler(BaseResourceHandler):
    """Handler for generating lesson plans as Word documents."""

    def __init__(self, structured_content: List[Dict[str, Any]], **kwargs):
        super().__init__(structured_content, **kwargs)
        if kwargs.get('include_images'):
            logger.info("Image support requested for lesson plan, but not implemented")

    def generate(self) -> str:
        """Generate the lesson plan docx file and return the file path."""
        # Create temp file
        temp_file = self.create_temp_file("docx")
        
        # Create a new document
        doc = docx.Document()
        
        # Add a title
        lesson_title = self.structured_content[0].get('title', 'Lesson Plan')
        lesson_title = self.clean_markdown_and_formatting(lesson_title)
        doc.add_heading(lesson_title, 0)
        
        # Add objectives section
        doc.add_heading('Learning Objectives', level=1)
        objectives = []
        for slide in self.structured_content:
            content = slide.get('content', [])
            for item in content:
                clean_item = self.clean_markdown_and_formatting(item)
                if 'objective' in clean_item.lower() or 'able to' in clean_item.lower():
                    objectives.append(clean_item)
        
        if objectives:
            for objective in objectives:
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(objective)
        else:
            doc.add_paragraph('Students will demonstrate understanding of the lesson content.')
                
        # Add materials section
        doc.add_heading('Materials', level=1)
        doc.add_paragraph('Standard classroom materials and any specific resources mentioned in the lesson content.')
            
        # Add procedure section with details from each slide
        doc.add_heading('Procedure', level=1)
        
        for i, slide in enumerate(self.structured_content):
            title = slide.get('title', f'Section {i+1}')
            clean_title = self.clean_markdown_and_formatting(title)
            doc.add_heading(clean_title, level=2)
            
            # Add content as procedure steps
            content = slide.get('content', [])
            clean_content = self.clean_content_list(content)
            if clean_content:
                for item in clean_content:
                    p = doc.add_paragraph()
                    p.add_run('• ').bold = True
                    p.add_run(item)
        
        # Add assessment section
        doc.add_heading('Assessment and Notes', level=1)
        doc.add_paragraph('Monitor student understanding through observation, questioning, and review of completed work.')
        doc.add_paragraph('Adjust pacing and provide additional support as needed based on student responses.')
        
        # Save the document
        doc.save(temp_file)
        
        # Verify file was created and is not empty
        if not os.path.exists(temp_file):
            raise FileNotFoundError(f"Failed to create lesson plan file at {temp_file}")
            
        file_size = os.path.getsize(temp_file)
        logger.info(f"Generated lesson plan file size: {file_size} bytes")
        
        if file_size == 0:
            raise ValueError("Generated lesson plan file is empty")
            
        return temp_file