# src/resource_handlers/quiz_handler.py - CLEANED VERSION
import os
import logging
import docx
import re
from typing import Dict, Any, List
from .base_handler import BaseResourceHandler

logger = logging.getLogger(__name__)

class QuizHandler(BaseResourceHandler):
    """Handler for generating quizzes as Word documents with multilingual support."""

    def __init__(self, structured_content: List[Dict[str, Any]], **kwargs):
        super().__init__(structured_content, **kwargs)
        if kwargs.get('include_images'):
            logger.info("Image support requested for quiz, but not implemented")

    def generate(self) -> str:
        """Generate a quiz docx file with properly separated questions and answers."""
        # Create temp file
        temp_file = self.create_temp_file("docx")
        
        # Create a new document
        doc = docx.Document()
        
        # Add title - use first section title or default
        quiz_title = "Quiz/Test"
        if self.structured_content and len(self.structured_content) > 0:
            raw_title = self.structured_content[0].get('title', 'Quiz/Test')
            quiz_title = self.clean_markdown_and_formatting(raw_title)
        
        doc.add_heading(quiz_title, 0)
        
        # Add name and date fields
        name_date_para = doc.add_paragraph()
        name_date_para.add_run('Name: ').bold = True
        name_date_para.add_run('_' * 40)
        name_date_para.add_run('    Date: ').bold = True
        name_date_para.add_run('_' * 20)
        
        doc.add_paragraph()  # Add spacing
        
        # STUDENT SECTION
        student_heading = doc.add_heading("STUDENT SECTION", 1)
        student_heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Add general instructions
        instructions_para = doc.add_paragraph()
        instructions_para.add_run('Instructions: ').bold = True
        instructions_para.add_run('Answer all questions to the best of your ability. Show your work where applicable.')
        
        doc.add_paragraph()  # Add some space
        
        # Process each section and extract questions
        question_counter = 1
        all_teacher_data = []  # Collect all teacher guidance for the answer key
        
        for section_idx, section in enumerate(self.structured_content):
            section_title = self.clean_markdown_and_formatting(section.get('title', f'Section {section_idx + 1}'))
            
            # Only process sections with structured questions - no legacy support
            if 'structured_questions' not in section or not section['structured_questions']:
                logger.warning(f"Skipping section {section_title} - no structured questions found")
                continue
                
            questions, answers = self._extract_from_structured_questions(section['structured_questions'])
            teacher_notes = section.get('teacher_notes', [])
            differentiation_tips = section.get('differentiation_tips', [])
            
            # Store teacher data for later use
            all_teacher_data.append({
                'section_title': section_title,
                'questions': questions.copy(),
                'answers': answers.copy(),
                'teacher_notes': teacher_notes.copy(),
                'differentiation_tips': differentiation_tips.copy(),
                'start_question_num': question_counter
            })
            
            if not questions:
                continue
            
            # Add section heading for students
            section_heading = doc.add_heading(section_title, level=2)
            
            # Add questions only (no answers, no teacher notes)
            for question in questions:
                # Check if it's a multiple choice question
                if re.search(r'\b[A-D]\)', question):
                    # Multiple choice - add question as is
                    question_para = doc.add_paragraph()
                    question_para.add_run(f"{question_counter}. ").bold = True
                    question_para.add_run(question)
                    
                    # Add some space for the answer
                    doc.add_paragraph("Answer: _____")
                else:
                    # Regular question - add question with answer space
                    question_para = doc.add_paragraph()
                    question_para.add_run(f"{question_counter}. ").bold = True
                    question_para.add_run(question)
                    
                    # Add answer space based on question type
                    if any(word in question.lower() for word in ['calculate', 'solve', 'find', 'what is', 'calcular', 'resolver', 'encontrar', 'qué es']):
                        # Math/calculation question - provide work space
                        doc.add_paragraph()
                        doc.add_paragraph("Show your work:")
                        doc.add_paragraph("_" * 60)
                        doc.add_paragraph("_" * 60)
                        doc.add_paragraph()
                        answer_para = doc.add_paragraph()
                        answer_para.add_run("Answer: ").bold = True
                        answer_para.add_run("_" * 30)
                    else:
                        # Short answer question
                        doc.add_paragraph("_" * 60)
                        doc.add_paragraph("_" * 60)
                
                doc.add_paragraph()  # Add spacing between questions
                question_counter += 1
        
        # TEACHER SECTION - Start on new page
        doc.add_page_break()
        teacher_heading = doc.add_heading("TEACHER GUIDE & ANSWER KEY", 1)
        teacher_heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Add teacher guidance and answers by section
        for section_data in all_teacher_data:
            if not section_data['questions']:  # Skip sections with no questions
                continue
            
            # Section heading
            doc.add_heading(section_data['section_title'], level=2)
            
            # Questions and Answers
            if section_data['questions'] and len(section_data['answers']) > 0:
                doc.add_heading("Questions & Answers", level=3)
                
                current_q_num = section_data['start_question_num']
                for i, question in enumerate(section_data['questions']):
                    # Question
                    q_para = doc.add_paragraph()
                    q_para.add_run(f"Q{current_q_num}: ").bold = True
                    q_para.add_run(question)
                    
                    # Answer
                    if i < len(section_data['answers']):
                        a_para = doc.add_paragraph()
                        a_para.add_run("Answer: ").bold = True
                        a_para.add_run(section_data['answers'][i])
                    else:
                        a_para = doc.add_paragraph()
                        a_para.add_run("Answer: ").bold = True
                        a_para.add_run("(Teacher to provide)")
                    
                    doc.add_paragraph()  # Spacing
                    current_q_num += 1
            
            # Teaching Notes
            if section_data['teacher_notes']:
                doc.add_heading("Teaching Notes", level=3)
                for note in section_data['teacher_notes']:
                    note_para = doc.add_paragraph()
                    note_para.add_run("• ").bold = True
                    note_para.add_run(note)
            
            # Differentiation Tips
            if section_data['differentiation_tips']:
                doc.add_heading("Differentiation Tips", level=3)
                for tip in section_data['differentiation_tips']:
                    tip_para = doc.add_paragraph()
                    tip_para.add_run("• ").bold = True
                    tip_para.add_run(tip)
            
            # Add spacing between sections
            doc.add_paragraph()
        
        # Save the document
        doc.save(temp_file)
        
        # Verify file was created
        if not os.path.exists(temp_file):
            raise FileNotFoundError(f"Failed to create quiz file at {temp_file}")
            
        file_size = os.path.getsize(temp_file)
        logger.info(f"Generated quiz file size: {file_size} bytes")
        
        if file_size == 0:
            raise ValueError("Generated quiz file is empty")
            
        return temp_file
    
    def _extract_from_structured_questions(self, structured_questions: List[Dict]) -> tuple[List[str], List[str]]:
        """Extract questions and answers from structured JSON format"""
        questions = []
        answers = []
        
        for q_data in structured_questions:
            question_text = q_data.get('question', '')
            answer = q_data.get('answer', '')
            explanation = q_data.get('explanation', '')
            q_type = q_data.get('type', 'short_answer')
            
            # Format question based on type
            if q_type == 'multiple_choice':
                options = q_data.get('options', [])
                if len(options) >= 4:
                    formatted_question = f"{question_text} A) {options[0]} B) {options[1]} C) {options[2]} D) {options[3]}"
                else:
                    formatted_question = question_text
            else:
                formatted_question = question_text
            
            # Create complete answer with explanation
            if explanation and explanation != answer:
                complete_answer = f"{answer} - {explanation}"
            else:
                complete_answer = answer
            
            questions.append(formatted_question)
            answers.append(complete_answer)
        
        logger.info(f"Extracted {len(questions)} questions from structured format")
        return questions, answers