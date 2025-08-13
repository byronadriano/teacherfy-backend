# src/resource_handlers/worksheet_handler.py - CLEANED VERSION
import os
import logging
import docx
from typing import Dict, Any, List
from .base_handler import BaseResourceHandler

logger = logging.getLogger(__name__)

class WorksheetHandler(BaseResourceHandler):
    """Handler for generating worksheets as Word documents with multilingual support."""
    
    def __init__(self, structured_content: List[Dict[str, Any]], **kwargs):
        super().__init__(structured_content, **kwargs)
        if kwargs.get('include_images'):
            logger.info("Image support requested for worksheet, but not implemented")
    
    def generate(self) -> str:
        """Generate a worksheet docx file with clean separation of student and teacher content."""
        # Create temp file
        temp_file = self.create_temp_file("docx")
        
        # Create a new document
        doc = docx.Document()
        
        # Add title - use first section title or default
        worksheet_title = "Worksheet"
        if self.structured_content and len(self.structured_content) > 0:
            raw_title = self.structured_content[0].get('title', 'Worksheet')
            worksheet_title = self.clean_markdown_and_formatting(raw_title)
        
        doc.add_heading(worksheet_title, 0)
        
        # Add name and date fields
        name_date_para = doc.add_paragraph()
        name_date_para.add_run('Name: ').bold = True
        name_date_para.add_run('_' * 40)
        name_date_para.add_run('    Date: ').bold = True
        name_date_para.add_run('_' * 20)
        
        doc.add_paragraph()  # Add spacing
        
        # STUDENT SECTION
        student_heading = doc.add_heading("STUDENT SECTION", 1)
        student_heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Add general instructions
        instructions_para = doc.add_paragraph()
        instructions_para.add_run('Instructions: ').bold = True
        instructions_para.add_run('Read each question carefully and provide your answer in the space provided.')
        
        doc.add_paragraph()  # Add some space
        
        # Process each section and extract questions
        question_counter = 1
        all_teacher_data = []  # Collect all teacher guidance for the answer key
        
        for section_idx, section in enumerate(self.structured_content):
            section_title = self.clean_markdown_and_formatting(section.get('title', f'Section {section_idx + 1}'))
            
            # Use structured questions if available (new format), fallback to legacy content parsing
            if 'structured_questions' in section and section['structured_questions']:
                questions, answers = self._extract_from_structured_questions(section['structured_questions'])
                teacher_notes = section.get('teacher_notes', [])
                differentiation_tips = section.get('differentiation_tips', [])
            else:
                # Legacy fallback: extract from content list
                if not section.get('content'):
                    continue
                    
                questions, answers = self.extract_questions_from_content(section.get('content', []))
                teacher_notes, differentiation_tips = self.extract_teacher_guidance(section.get('content', []))
                
                # Add any additional notes from the content itself
                content_items = section.get('content', [])
                for item in content_items:
                    clean_item = self.clean_markdown_and_formatting(item)
                    # Look for teacher guidance within the content
                    if any(keyword in clean_item.lower() for keyword in ['teacher note:', 'differentiation tip:']):
                        if 'differentiation tip:' in clean_item.lower():
                            tip = clean_item.split('differentiation tip:', 1)[1].strip()
                            if tip:
                                differentiation_tips.append(tip)
                        elif 'teacher note:' in clean_item.lower():
                            note = clean_item.split('teacher note:', 1)[1].strip()
                            if note:
                                teacher_notes.append(note)
            
            # Store teacher data for later use
            all_teacher_data.append({
                'section_title': section_title,
                'questions': questions.copy(),
                'answers': answers.copy(),
                'teacher_notes': teacher_notes.copy(),
                'differentiation_tips': differentiation_tips.copy(),
                'start_question_num': question_counter
            })
            
            if not questions:
                continue
            
            # Add section heading for students
            section_heading = doc.add_heading(section_title, level=2)
            
            # Add questions only (no answers, no teacher notes)
            for question in questions:
                # Clean the question text to avoid Word document corruption
                clean_question = self.clean_markdown_and_formatting(question)
                # Remove any problematic characters that might cause corruption
                clean_question = clean_question.encode('ascii', errors='ignore').decode('ascii')
                
                # Check if it's a multiple choice question (contains newlines with A), B), etc.)
                if '\n' in clean_question and re.search(r'\b[A-D]\)', clean_question):
                    # Multiple choice - split into question and options
                    lines = clean_question.split('\n')
                    main_question = lines[0]
                    
                    # Add main question
                    question_para = doc.add_paragraph()
                    question_para.add_run(f"{question_counter}. ").bold = True
                    question_para.add_run(main_question)
                    
                    # Add each option on separate line with proper spacing
                    for line in lines[1:]:
                        if line.strip():
                            option_para = doc.add_paragraph()
                            option_para.add_run(f"   {line.strip()}")
                    
                    # Add answer space
                    doc.add_paragraph()
                    doc.add_paragraph("Answer: " + "_" * 20)
                else:
                    # Regular question
                    question_para = doc.add_paragraph()
                    question_para.add_run(f"{question_counter}. ").bold = True
                    question_para.add_run(clean_question)
                    
                    # Add answer space
                    if question.endswith('?') or '?' in question:
                        # For questions, add a line for answers
                        doc.add_paragraph("Answer: " + "_" * 50)
                    else:
                        # For fill-in-the-blank or calculation problems
                        doc.add_paragraph("_" * 60)
                
                doc.add_paragraph()  # Add spacing between questions
                question_counter += 1
        
        # TEACHER SECTION - Start on new page
        doc.add_page_break()
        teacher_heading = doc.add_heading("TEACHER GUIDE & ANSWER KEY", 1)
        teacher_heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Add teacher guidance and answers by section
        for section_data in all_teacher_data:
            if not section_data['questions']:  # Skip sections with no questions
                continue
            
            # Section heading
            doc.add_heading(section_data['section_title'], level=2)
            
            # Questions and Answers
            if section_data['questions'] and section_data['answers']:
                doc.add_heading("Questions & Answers", level=3)
                
                current_q_num = section_data['start_question_num']
                for i, question in enumerate(section_data['questions']):
                    # Question
                    q_para = doc.add_paragraph()
                    q_para.add_run(f"Q{current_q_num}: ").bold = True
                    q_para.add_run(question)
                    
                    # Answer
                    if i < len(section_data['answers']):
                        a_para = doc.add_paragraph()
                        a_para.add_run("Answer: ").bold = True
                        a_para.add_run(section_data['answers'][i])
                    else:
                        a_para = doc.add_paragraph()
                        a_para.add_run("Answer: ").bold = True
                        a_para.add_run("(Teacher to provide)")
                    
                    doc.add_paragraph()  # Spacing
                    current_q_num += 1
            
            # Teaching Notes
            if section_data['teacher_notes']:
                doc.add_heading("Teaching Notes", level=3)
                for note in section_data['teacher_notes']:
                    note_para = doc.add_paragraph()
                    note_para.add_run("• ").bold = True
                    note_para.add_run(note)
            
            # Differentiation Tips
            if section_data['differentiation_tips']:
                doc.add_heading("Differentiation Tips", level=3)
                for tip in section_data['differentiation_tips']:
                    tip_para = doc.add_paragraph()
                    tip_para.add_run("• ").bold = True
                    tip_para.add_run(tip)
            
            # Add spacing between sections
            doc.add_paragraph()
        
        # Save the document
        doc.save(temp_file)
        
        # Verify file was created
        if not os.path.exists(temp_file):
            raise FileNotFoundError(f"Failed to create worksheet file at {temp_file}")
            
        file_size = os.path.getsize(temp_file)
        logger.info(f"Generated worksheet file size: {file_size} bytes")
        
        if file_size == 0:
            raise ValueError("Generated worksheet file is empty")
            
        return temp_file

    def _extract_from_structured_questions(self, structured_questions: List[Dict]) -> tuple[List[str], List[str]]:
        """Extract questions and answers from structured JSON format with improved formatting"""
        questions = []
        answers = []
        
        for q_data in structured_questions:
            question_text = q_data.get('question', '')
            answer = q_data.get('answer', '')
            explanation = q_data.get('explanation', '')
            teacher_instruction = q_data.get('teacher_instruction', '')
            q_type = q_data.get('type', 'short_answer')
            
            # Format question based on type
            if q_type == 'multiple_choice':
                options = q_data.get('options', [])
                if len(options) >= 4:
                    # Format with newlines for proper separation in document
                    formatted_question = f"{question_text}\nA) {options[0]}\nB) {options[1]}\nC) {options[2]}\nD) {options[3]}"
                else:
                    formatted_question = question_text
            elif q_type == 'fill_blank':
                # Ensure fill-in-the-blank format
                if '___' not in question_text and '_' not in question_text:
                    formatted_question = f"{question_text} ___"
                else:
                    formatted_question = question_text
            else:
                formatted_question = question_text
            
            # Create complete answer with explanation and teacher instruction for teacher guide
            complete_answer = answer
            if explanation and explanation != answer:
                complete_answer += f" ({explanation})"
            if teacher_instruction:
                complete_answer += f" [Teacher instruction: {teacher_instruction}]"
            
            questions.append(formatted_question)
            answers.append(complete_answer)
        
        logger.info(f"Extracted {len(questions)} worksheet questions from structured format")
        return questions, answers