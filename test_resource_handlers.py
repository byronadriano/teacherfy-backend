# test_resource_handlers.py
import json
import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.resource_handlers.worksheet_handler import WorksheetHandler
from src.resource_handlers.quiz_handler import QuizHandler

# Test data that simulates the problematic structure
test_worksheet_data = [
    {
        "title": "Práctica de Suma y Resta",
        "layout": "TITLE_AND_CONTENT",
        "content": [
            "Resuelve: 123 + 456 = ____ (Answer: 579)",
            "Resuelve: 789 - 432 = ____ (Answer: 357)",
            "Encuentra la suma de 234 y 567. ¿Cuál es el resultado? (Answer: 801)",
            "Resta 128 a 654. ¿Cuánto queda? (Answer: 526)",
            "Differentiation tip: Para estudiantes que necesiten apoyo, utiliza bloques base diez para representar visualmente las sumas y restas.",
            "Teacher note: Asegúrese de repasar cómo alinear los números correctamente antes de comenzar las operaciones."
        ]
    },
    {
        "title": "Multiplicación Básica",
        "layout": "TITLE_AND_CONTENT",
        "content": [
            "Multiplica: 12 x 3 = ____ (Answer: 36)",
            "Multiplica: 25 x 4 = ____ (Answer: 100)",
            "¿Cuál es el producto de 7 x 8? (Answer: 56)",
            "Encuentra el resultado de 9 x 6. (Answer: 54)",
            "Differentiation tip: Introduce la multiplicación usando arrays o grupos de objetos para visualizar los problemas.",
            "Teacher note: Anime a los estudiantes a practicar las tablas de multiplicar regularmente para mejorar su fluidez."
        ]
    },
    {
        "title": "Resolución de Problemas con Palabras",
        "layout": "TITLE_AND_CONTENT",
        "content": [
            "Marta tiene 3 bolsas con 12 manzanas cada una. ¿Cuántas manzanas tiene en total? (Answer: 36)",
            "Si un tren viaja a 60 kilómetros por hora y viaja por 2 horas, ¿cuántos kilómetros viaja en total? (Answer: 120)",
            "Juan reparte 48 galletas en paquetes de 6. ¿Cuántos paquetes hace? (Answer: 8)",
            "Ana compra 4 paquetes de stickers, cada paquete tiene 9 stickers. ¿Cuántos stickers compra en total? (Answer: 36)",
            "Differentiation tip: Para estudiantes avanzados, pide que escriban sus propios problemas de palabras basados en situaciones de la vida real.",
            "Teacher note: Discutir las estrategias para resolver problemas de palabras, como subrayar la información clave y dibujar un boceto."
        ]
    }
]

test_quiz_data = [
    {
        "title": "Understanding Fractions",
        "layout": "TITLE_AND_CONTENT",
        "content": [
            "What fraction represents half of a pizza? A) 1/4 B) 1/2 C) 3/4 D) 1/3 (Answer: B)",
            "Write the fraction that shows 3 out of 8 equal parts. (Answer: 3/8)",
            "Which is larger: 2/5 or 3/7? (Answer: 3/7)",
            "Differentiation tip: Allow students to sketch visuals or use manipulatives",
            "Teacher note: Review numerator and denominator concepts before starting"
        ]
    },
    {
        "title": "Fraction Operations",
        "layout": "TITLE_AND_CONTENT", 
        "content": [
            "Add: 1/4 + 2/4 = ? (Answer: 3/4)",
            "Subtract: 5/6 - 2/6 = ? (Answer: 3/6 or 1/2)",
            "Convert 2 1/3 to an improper fraction. (Answer: 7/3)",
            "Differentiation tip: Use fraction strips for visual support",
            "Teacher note: Emphasize that denominators must be the same when adding/subtracting"
        ]
    }
]

def test_worksheet_handler():
    """Test the worksheet handler with the problematic data"""
    print("=== Testing Worksheet Handler ===")
    
    try:
        handler = WorksheetHandler(test_worksheet_data)
        output_file = handler.generate()
        
        print(f"✅ Worksheet generated successfully: {output_file}")
        print(f"File size: {os.path.getsize(output_file)} bytes")
        
        # Verify file exists and has content
        if os.path.exists(output_file) and os.path.getsize(output_file) > 1000:
            print("✅ Worksheet file appears to be properly generated")
            
            # Try to read the document to verify structure
            try:
                import docx
                doc = docx.Document(output_file)
                
                print(f"Document has {len(doc.paragraphs)} paragraphs")
                
                # Look for key sections
                student_section_found = False
                teacher_section_found = False
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if "STUDENT SECTION" in text:
                        student_section_found = True
                    elif "TEACHER GUIDE" in text:
                        teacher_section_found = True
                
                if student_section_found and teacher_section_found:
                    print("✅ Both student and teacher sections found")
                else:
                    print(f"❌ Missing sections - Student: {student_section_found}, Teacher: {teacher_section_found}")
                
                # Check for answer separation
                mixed_answers = 0
                for para in doc.paragraphs:
                    if "(Answer:" in para.text and not "TEACHER" in para.text:
                        mixed_answers += 1
                
                if mixed_answers > 0:
                    print(f"❌ Found {mixed_answers} mixed answers in student section")
                else:
                    print("✅ No mixed answers found in student section")
                    
            except Exception as e:
                print(f"⚠️ Could not analyze document structure: {e}")
        else:
            print("❌ Worksheet file is missing or too small")
            
    except Exception as e:
        print(f"❌ Worksheet generation failed: {e}")
        import traceback
        traceback.print_exc()

def test_quiz_handler():
    """Test the quiz handler"""
    print("\n=== Testing Quiz Handler ===")
    
    try:
        handler = QuizHandler(test_quiz_data)
        output_file = handler.generate()
        
        print(f"✅ Quiz generated successfully: {output_file}")
        print(f"File size: {os.path.getsize(output_file)} bytes")
        
        # Verify file exists and has content
        if os.path.exists(output_file) and os.path.getsize(output_file) > 1000:
            print("✅ Quiz file appears to be properly generated")
            
            # Try to read the document to verify structure
            try:
                import docx
                doc = docx.Document(output_file)
                
                print(f"Document has {len(doc.paragraphs)} paragraphs")
                
                # Look for key sections
                student_section_found = False
                teacher_section_found = False
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if "STUDENT SECTION" in text:
                        student_section_found = True
                    elif "TEACHER GUIDE" in text:
                        teacher_section_found = True
                
                if student_section_found and teacher_section_found:
                    print("✅ Both student and teacher sections found")
                else:
                    print(f"❌ Missing sections - Student: {student_section_found}, Teacher: {teacher_section_found}")
                
                # Check for answer separation
                mixed_answers = 0
                for para in doc.paragraphs:
                    if "(Answer:" in para.text and not "TEACHER" in para.text:
                        mixed_answers += 1
                
                if mixed_answers > 0:
                    print(f"❌ Found {mixed_answers} mixed answers in student section")
                else:
                    print("✅ No mixed answers found in student section")
                    
            except Exception as e:
                print(f"⚠️ Could not analyze document structure: {e}")
        else:
            print("❌ Quiz file is missing or too small")
            
    except Exception as e:
        print(f"❌ Quiz generation failed: {e}")
        import traceback
        traceback.print_exc()

def test_content_parsing():
    """Test the content parsing functions directly"""
    print("\n=== Testing Content Parsing Functions ===")
    
    from src.resource_handlers.worksheet_handler import extract_questions_from_content, extract_teacher_guidance
    
    # Test content with mixed elements
    test_content = [
        "Solve: 25 + 17 = ____ (Answer: 42)",
        "What is 34 + 28? (Answer: 62)",
        "Differentiation tip: Use number lines for visual learners",
        "Teacher note: Remind students about place value",
        "Fill in: 45 + ___ = 73 (Answer: 28)"
    ]
    
    questions, answers = extract_questions_from_content(test_content)
    teacher_notes, diff_tips = extract_teacher_guidance(test_content)
    
    print(f"Questions extracted: {len(questions)}")
    for i, q in enumerate(questions):
        print(f"  Q{i+1}: {q}")
    
    print(f"Answers extracted: {len(answers)}")
    for i, a in enumerate(answers):
        print(f"  A{i+1}: {a}")
    
    print(f"Teacher notes: {len(teacher_notes)}")
    for note in teacher_notes:
        print(f"  Note: {note}")
    
    print(f"Differentiation tips: {len(diff_tips)}")
    for tip in diff_tips:
        print(f"  Tip: {tip}")
    
    # Verify separation worked
    if len(questions) == 3 and len(answers) == 3:
        print("✅ Question/answer extraction working correctly")
    else:
        print(f"❌ Expected 3 questions and 3 answers, got {len(questions)} questions and {len(answers)} answers")
    
    if len(teacher_notes) == 1 and len(diff_tips) == 1:
        print("✅ Teacher guidance extraction working correctly")
    else:
        print(f"❌ Expected 1 teacher note and 1 tip, got {len(teacher_notes)} notes and {len(diff_tips)} tips")

if __name__ == "__main__":
    print("Testing Resource Handler Fixes")
    print("=" * 50)
    
    test_content_parsing()
    test_worksheet_handler()
    test_quiz_handler()
    
    print("\n" + "=" * 50)
    print("Testing completed!")