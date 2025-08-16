# resources/handlers/lesson_plan_handler.py - CLEANED VERSION
import os
import logging
import docx
from typing import Dict, Any, List
from .base_handler import BaseResourceHandler

logger = logging.getLogger(__name__)

class LessonPlanHandler(BaseResourceHandler):
    """Handler for generating lesson plans as Word documents."""

    def __init__(self, structured_content: List[Dict[str, Any]], **kwargs):
        super().__init__(structured_content, **kwargs)
        if kwargs.get('include_images'):
            logger.info("Image support requested for lesson plan, but not implemented")

    def generate(self) -> str:
        """Generate the lesson plan docx file and return the file path."""
        # Create temp file
        temp_file = self.create_temp_file("docx")
        
        # Create a new document
        doc = docx.Document()
        
        # Add a title
        lesson_title = self.structured_content[0].get('title', 'Lesson Plan')
        lesson_title = self.clean_markdown_and_formatting(lesson_title)
        doc.add_heading(lesson_title, 0)
        
        # Add objectives section
        doc.add_heading('Learning Objectives', level=1)
        objectives = []
        all_materials = set()
        
        # Extract objectives and materials from structured content
        for section in self.structured_content:
            # Look for objectives in structured activities
            if 'structured_activities' in section:
                for activity in section['structured_activities']:
                    activity_text = activity.get('activity', '')
                    if 'objective' in activity_text.lower() or 'students will' in activity_text.lower():
                        objectives.append(activity_text)
                    # Collect materials
                    materials = activity.get('materials', [])
                    all_materials.update(materials)
            
            # Also check legacy content format
            content = section.get('content', [])
            for item in content:
                clean_item = self.clean_markdown_and_formatting(item)
                if 'objective' in clean_item.lower() or 'able to' in clean_item.lower():
                    objectives.append(clean_item)
        
        # Add objectives or generate from lesson title
        if objectives:
            for objective in objectives:
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(objective)
        else:
            # Generate basic objective from lesson title
            p = doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(f'Students will understand key concepts related to {lesson_title.lower()}')
            p = doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(f'Students will be able to apply {lesson_title.lower()} in various contexts')
                
        # Add materials section
        doc.add_heading('Materials', level=1)
        if all_materials:
            for material in sorted(all_materials):
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(material)
        else:
            doc.add_paragraph('Standard classroom materials and any specific resources mentioned in the lesson content.')
            
        # Add procedure section with details from each slide
        doc.add_heading('Procedure', level=1)
        
        for i, section in enumerate(self.structured_content):
            title = section.get('title', f'Section {i+1}')
            clean_title = self.clean_markdown_and_formatting(title)
            doc.add_heading(clean_title, level=2)
            
            # Use structured activities if available (new format), fallback to legacy content
            if 'structured_activities' in section and section['structured_activities']:
                activities = section['structured_activities']
                teacher_actions = section.get('teacher_actions', [])
                differentiation_tips = section.get('differentiation_tips', [])
                assessment_checks = section.get('assessment_checks', [])
                
                # Add activities
                for activity_data in activities:
                    activity_text = activity_data.get('activity', '')
                    duration = activity_data.get('duration', '')
                    materials = activity_data.get('materials', [])
                    instructions = activity_data.get('instructions', '')
                    
                    # Activity description
                    p = doc.add_paragraph()
                    p.add_run('Activity: ').bold = True
                    p.add_run(f"{activity_text}")
                    if duration:
                        p.add_run(f" ({duration})")
                    
                    # Materials
                    if materials:
                        materials_p = doc.add_paragraph()
                        materials_p.add_run('Materials: ').bold = True
                        materials_p.add_run(', '.join(materials))
                    
                    # Instructions
                    if instructions:
                        inst_p = doc.add_paragraph()
                        inst_p.add_run('Instructions: ').bold = True
                        inst_p.add_run(instructions)
                    
                    doc.add_paragraph()  # Add spacing
                
                # Add teacher actions
                if teacher_actions:
                    ta_heading = doc.add_heading('Teacher Actions', level=3)
                    for action in teacher_actions:
                        p = doc.add_paragraph()
                        p.add_run('• ').bold = True
                        p.add_run(action)
                
                # Add differentiation tips
                if differentiation_tips:
                    diff_heading = doc.add_heading('Differentiation Tips', level=3)
                    for tip in differentiation_tips:
                        p = doc.add_paragraph()
                        p.add_run('• ').bold = True
                        p.add_run(tip)
                
                # Add assessment checks
                if assessment_checks:
                    assess_heading = doc.add_heading('Assessment Checks', level=3)
                    for check in assessment_checks:
                        p = doc.add_paragraph()
                        p.add_run('• ').bold = True
                        p.add_run(check)
                        
            else:
                # Legacy fallback: use content list
                content = section.get('content', [])
                clean_content = self.clean_content_list(content)
                if clean_content:
                    for item in clean_content:
                        p = doc.add_paragraph()
                        p.add_run('• ').bold = True
                        p.add_run(item)
        
        # Add assessment section
        doc.add_heading('Assessment and Notes', level=1)
        doc.add_paragraph('Monitor student understanding through observation, questioning, and review of completed work.')
        doc.add_paragraph('Adjust pacing and provide additional support as needed based on student responses.')
        
        # Save the document
        doc.save(temp_file)
        
        # Verify file was created and is not empty
        if not os.path.exists(temp_file):
            raise FileNotFoundError(f"Failed to create lesson plan file at {temp_file}")
            
        file_size = os.path.getsize(temp_file)
        logger.info(f"Generated lesson plan file size: {file_size} bytes")
        
        if file_size == 0:
            raise ValueError("Generated lesson plan file is empty")
            
        return temp_file